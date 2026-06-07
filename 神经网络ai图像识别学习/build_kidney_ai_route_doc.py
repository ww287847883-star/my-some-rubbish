from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\14901\Desktop\数学，统计建模培训\神经网络ai图像识别学习\基于人工智能的肾肿瘤侵袭性与惰性预测系统_学习与制作路线.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Microsoft YaHei"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Microsoft YaHei"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    else:
        r = p.add_run(text)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "1F4E79")
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
        for c in cells:
            set_cell_shading(c, "FFFFFF")
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for m in ["top", "left", "bottom", "right"]:
                node = OxmlElement(f"w:{m}")
                node.set(qn("w:w"), "90")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)
    doc.add_paragraph()
    return table


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.25)
    sec.right_margin = Cm(2.25)
    styles = doc.styles
    for name in ["Normal", "Body Text"]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 20, "17365D"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12.5, "2F5597"),
        ("Heading 3", 11.5, "365F91"),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True


def add_callout(doc, title, body, fill="EAF3F8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10)
    doc.add_paragraph()


doc = Document()
style_doc(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("基于人工智能的肾肿瘤侵袭性与惰性预测系统研究")
r.bold = True
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(23, 54, 93)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("零基础学习路径与可落地制作路线脚本")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(89, 89, 89)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("参考材料：冠脉DSA影像狭窄自动检测与分割系统研究申请书；文献9《Artificial intelligence links CT images to pathologic features and survival outcomes of renal masses》")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(100, 100, 100)

add_callout(
    doc,
    "一句话目标",
    "先做一个可运行、可展示、可写进论文的轻量系统：输入肾肿瘤 CT 图像或肿瘤 ROI，完成预处理、模型预测、结果可视化和报告输出；模型目标是区分“侵袭性”与“惰性”，后续再扩展到自动分割、多期相融合和生存风险分析。",
    "EAF3F8",
)

doc.add_heading("一、把课题降级成可完成的 MVP", level=1)
add_paragraph(
    doc,
    "你的题目很专业，完整版本会涉及多期相 CT、肿瘤分割、病理标签、深度学习分类、可解释性和软件系统。零基础阶段不建议一开始追求顶级论文复现，而应先做最小可用系统。",
)
add_table(
    doc,
    ["版本", "目标", "输入", "模型", "输出", "难度"],
    [
        ["MVP 1", "肿瘤侵袭性/惰性二分类", "医生/人工裁剪后的肾肿瘤 ROI 图像", "ResNet18 或 2.5D CNN", "概率、类别、Grad-CAM 热力图", "低"],
        ["MVP 2", "加入 CT 医学预处理", "DICOM/NIfTI CT 切片", "ResNet18 + MONAI transforms", "预测结果和病例记录", "中"],
        ["MVP 3", "半自动分割 + 分类", "CT 体数据 + 肿瘤 mask", "U-Net/nnU-Net + 分类模型", "肿瘤区域、侵袭性概率", "中高"],
        ["论文增强版", "多期相 CT 融合预测", "平扫/动脉期/静脉期 CT", "多分支 CNN 或 3D CNN", "AUC、敏感度、特异度、可解释性", "高"],
    ],
    widths=[2.0, 3.2, 3.4, 3.0, 3.4, 1.5],
)

add_callout(
    doc,
    "建议选择",
    "先完成 MVP 1 + MVP 2。它不算高级，但完整可用：有数据处理、有模型训练、有验证指标、有界面展示，足够支撑课程项目、开题、大创初稿和后续论文雏形。",
    "FFF2CC",
)

doc.add_heading("二、两篇参考材料给你的启发", level=1)
add_table(
    doc,
    ["参考材料", "可借鉴内容", "迁移到本课题的做法"],
    [
        [
            "冠脉 DSA 系统申请书",
            "写作结构清晰：研究目的、研究内容、国内外现状、创新点、技术路线、软件模块、进度安排。",
            "你的文档也按“模型开发 + 软件系统”两条线展开，系统模块可设置为图像导入、预处理、模型推理、结果分析、可视化报告。",
        ],
        [
            "文献9 肾肿瘤 CT AI",
            "使用术前 CT 预测肾肿瘤良恶性和侵袭性；大样本、多中心、多期相 CT；分割网络先定位肿瘤，再用多期相 CNN 分类。",
            "初期不用复现全部规模，先采用 ROI 分类。论文写法上强调非侵入式预测、病理标签、AUC 评价、与传统影像组学或人工读片对比。",
        ],
    ],
    widths=[3.2, 5.2, 6.2],
)

doc.add_heading("三、文献9的核心技术路线总结", level=1)
add_bullets(
    doc,
    [
        "研究问题：术前 CT 图像能否预测肾肿瘤的病理特征，并进一步区分侵袭性与惰性。",
        "数据规模：文献纳入 4557 名患者、13261 个术前 CT 体数据，包含训练集、内部测试集、外部测试集、前瞻性测试集和 TCIA 测试集。",
        "模型框架：先用分割网络描绘肾肿瘤区域，分割性能 DICE 约 0.852；再基于裁剪后的肿瘤图像建立多期相卷积神经网络。",
        "任务结果：良恶性预测在前瞻性测试集中 AUC 约 0.871；侵袭性与惰性预测在前瞻性测试集中 AUC 约 0.783。",
        "可解释性：使用 CAM/类似热力图展示模型关注区域，帮助说明模型不是只输出黑盒概率。",
        "临床意义：模型可作为术前辅助决策工具，减少对病理不确定性的依赖，辅助选择随访、消融或手术治疗策略。",
    ]
)

doc.add_heading("四、你的系统推荐技术路线", level=1)
add_paragraph(doc, "建议采用“先分类、后分割、再多期相”的路线，避免一开始被医学影像三维数据和复杂标注卡住。")
add_table(
    doc,
    ["步骤", "具体任务", "产物"],
    [
        ["1", "收集肾肿瘤 CT 图像及病理标签，至少明确每例是侵袭性还是惰性。", "病例表、标签表、数据字典"],
        ["2", "用 3D Slicer 或 ITK-SNAP 手动裁剪/标注肿瘤 ROI。", "ROI 图像或 mask"],
        ["3", "完成 CT 预处理：读取 DICOM/NIfTI、窗宽窗位、归一化、resize、数据增强。", "preprocess.py"],
        ["4", "训练一个基础分类模型：ResNet18、DenseNet121 或简单 3D CNN。", "train.py、best_model.pth"],
        ["5", "评价模型：AUC、Accuracy、Sensitivity、Specificity、Confusion Matrix。", "metrics.csv、ROC 曲线"],
        ["6", "加入 Grad-CAM，显示模型关注的肿瘤区域。", "热力图输出"],
        ["7", "做一个简易界面：上传图像、点击预测、显示结果、保存报告。", "Streamlit/PyQt 系统原型"],
    ],
    widths=[1.2, 9.2, 4.2],
)

doc.add_heading("五、零基础学习路径", level=1)
add_table(
    doc,
    ["阶段", "学习内容", "必须会做的操作", "建议时长"],
    [
        ["1", "Python + NumPy", "读写文件、数组切片、reshape、保存结果", "2-3 周"],
        ["2", "医学图像基础", "理解 DICOM、NIfTI、HU 值、窗宽窗位、slice/voxel", "2 周"],
        ["3", "PyTorch 基础", "Dataset、DataLoader、模型、loss、optimizer、训练循环", "3-4 周"],
        ["4", "CNN 图像分类", "训练 ResNet18 做二分类，画 ROC 和混淆矩阵", "3 周"],
        ["5", "CT 项目实战", "处理肾肿瘤 ROI，完成侵袭性/惰性预测", "4-6 周"],
        ["6", "可视化与系统", "Grad-CAM、Streamlit/PyQt、结果报告", "2-3 周"],
        ["7", "论文与答辩", "写技术路线、实验设计、结果分析、创新点", "2 周"],
    ],
    widths=[1.5, 3.3, 7.0, 2.0],
)

doc.add_heading("六、项目文件夹与脚本设计", level=1)
add_paragraph(doc, "建议从一开始就按工程项目管理，不要把数据、代码、模型和结果混在一个文件夹里。")
add_table(
    doc,
    ["目录/脚本", "作用"],
    [
        ["data/raw", "原始 DICOM、NIfTI 或导出的 CT 图片，只读保存。"],
        ["data/roi", "裁剪后的肾肿瘤 ROI 图像或 npy 文件。"],
        ["data/labels.csv", "病例编号、标签、期相、数据划分等信息。"],
        ["src/preprocess.py", "读取医学影像、窗宽窗位、归一化、resize。"],
        ["src/dataset.py", "定义 PyTorch Dataset 和数据增强。"],
        ["src/model.py", "定义 ResNet18/3D CNN 模型。"],
        ["src/train.py", "训练模型并保存 best_model.pth。"],
        ["src/evaluate.py", "计算 AUC、敏感度、特异度、混淆矩阵。"],
        ["src/gradcam.py", "生成模型热力图。"],
        ["app/app.py", "简易预测系统界面。"],
        ["outputs", "保存模型、图表、日志和预测报告。"],
    ],
    widths=[4.0, 10.4],
)

doc.add_heading("七、核心训练脚本伪代码", level=1)
add_paragraph(doc, "下面不是最终代码，而是你后续写程序时要遵循的脚本骨架。")
code = [
    "读取 labels.csv，按患者级别划分 train/val/test",
    "Dataset 根据病例路径读取 CT ROI 或切片图像",
    "对图像执行窗宽窗位、归一化、resize、随机翻转/旋转",
    "加载 ResNet18，将最后一层改为二分类输出",
    "使用 BCEWithLogitsLoss 或 CrossEntropyLoss",
    "每个 epoch 训练后在验证集计算 AUC",
    "保存验证集 AUC 最高的模型 best_model.pth",
    "在测试集输出 AUC、Accuracy、Sensitivity、Specificity",
    "对典型病例生成 Grad-CAM 热力图",
]
add_numbered(doc, code)

doc.add_heading("八、系统界面功能设计", level=1)
add_table(
    doc,
    ["模块", "功能", "初版实现"],
    [
        ["图像导入", "选择 CT 图片、ROI、DICOM 或 NIfTI 文件。", "先支持 PNG/JPG/NPY，后扩展 DICOM。"],
        ["预处理", "统一图像尺寸、灰度归一化、窗宽窗位。", "封装为 preprocess_image()。"],
        ["模型预测", "加载训练好的模型，输出侵袭性概率。", "PyTorch 推理。"],
        ["结果分析", "显示类别、概率、风险提示、模型置信度。", "表格 + 文本结果。"],
        ["可解释性", "叠加 Grad-CAM 热力图。", "保存 overlay 图片。"],
        ["报告导出", "导出病例编号、预测结果、热力图。", "先导出 PNG/CSV，后续导出 Word/PDF。"],
    ],
    widths=[2.2, 5.8, 6.2],
)

doc.add_heading("九、论文/申报书写作框架", level=1)
add_bullets(
    doc,
    [
        "研究目的：利用人工智能从术前 CT 中预测肾肿瘤侵袭性与惰性，为治疗方案选择提供辅助依据。",
        "研究内容：构建数据集；完成 CT 预处理；训练分类模型；设计可解释性模块；开发简易预测系统。",
        "拟解决问题：肾肿瘤术前病理不确定、人工判断主观性强、传统影像特征难以全面表达肿瘤异质性。",
        "创新点：把肾肿瘤 CT 影像、深度学习分类、可解释热力图和软件系统集成到一个完整流程中。",
        "预期成果：训练得到一个侵袭性/惰性预测模型，完成一个可演示软件原型，形成实验结果图表和项目论文初稿。",
    ]
)

doc.add_heading("十、3个月可执行进度表", level=1)
add_table(
    doc,
    ["时间", "任务", "验收标准"],
    [
        ["第1-2周", "学习 Python、NumPy、医学图像基本概念。", "能读取图像并显示 CT 灰度图。"],
        ["第3-4周", "学习 PyTorch 分类流程，先跑通猫狗/普通医学图片二分类。", "能完成 train/val/test 训练流程。"],
        ["第5-6周", "整理肾肿瘤数据和标签，制作 ROI 数据集。", "形成 labels.csv 和 data/roi。"],
        ["第7-8周", "训练 ResNet18/DenseNet121 基线模型。", "输出 AUC、混淆矩阵、ROC 曲线。"],
        ["第9-10周", "加入 Grad-CAM 和错误样本分析。", "生成可解释性图片。"],
        ["第11-12周", "完成简易软件界面和报告输出。", "能上传图像并显示预测结果。"],
    ],
    widths=[2.2, 8.0, 4.2],
)

doc.add_heading("十一、风险与降级方案", level=1)
add_table(
    doc,
    ["风险", "表现", "降级方案"],
    [
        ["数据太少", "模型过拟合，验证集不稳定。", "先做方法流程演示，使用交叉验证和数据增强，明确写为探索性研究。"],
        ["没有分割标注", "无法训练 U-Net。", "先用医生/人工裁剪 ROI，分割作为后续工作。"],
        ["DICOM 难处理", "读取方向、窗宽窗位混乱。", "先将 CT 导出为 PNG/NPY，之后再补 DICOM 读取。"],
        ["电脑算力不足", "3D 模型训练慢。", "先用 2D 或 2.5D 模型，只取肿瘤最大层面及相邻切片。"],
        ["结果不够高", "AUC 不理想。", "重点展示完整系统流程、可解释性和误差分析，不夸大临床能力。"],
    ],
    widths=[2.4, 4.4, 7.6],
)

doc.add_heading("十二、结论：最适合你的路线", level=1)
add_paragraph(
    doc,
    "最稳妥的路线是：先用裁剪后的肾肿瘤 ROI 做侵袭性/惰性二分类，跑通 PyTorch 训练、评估和 Grad-CAM；再把模型嵌入一个 Streamlit 或 PyQt 小系统；最后在论文中把它表述为“基于 CT 影像的肾肿瘤侵袭性预测系统原型”。这样技术难度可控，成果又完整。",
)
add_paragraph(
    doc,
    "后续升级方向是自动分割、多期相 CT 融合、3D CNN/MONAI、nnU-Net 和与病理/临床指标联合建模。不要一开始就追求文献9的完整规模，先做出能运行、能解释、能展示的基础系统。",
)

doc.add_heading("参考文献与材料来源", level=1)
add_bullets(
    doc,
    [
        "冠脉DSA影像狭窄自动检测与分割系统研究(1)——参考.doc：用于借鉴课题申报书结构、技术路线图思路和软件模块划分。",
        "Ying Xiong et al. Artificial intelligence links CT images to pathologic features and survival outcomes of renal masses. Nature Communications, 2025. DOI: 10.1038/s41467-025-56784-z。",
    ]
)

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("肾肿瘤侵袭性与惰性预测系统学习与制作路线")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)

for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

doc.save(OUT)
print(OUT)
